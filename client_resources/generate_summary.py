import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import glob, re, os

CONSULTATIONS_DIR = r"C:\Users\twpot\Documents\Tyler's Stuff\Tyler's Documents\Investing\Freelancing\Consultations"
CONSULTATION_NUMBER = 149

# Use the most recent consultation docx as a template so styles (especially
# List Paragraph with bullet numbering) are inherited correctly.
# A fresh docx.Document() lacks the bullet numbering definition.
template_files = sorted(
    [f for f in glob.glob(os.path.join(CONSULTATIONS_DIR, 'Consultation*.docx'))
     if re.search(r'Consultation\d+\.docx$', os.path.basename(f))
     and int(re.search(r'\d+', os.path.basename(f)).group()) != CONSULTATION_NUMBER],
    key=lambda f: int(re.search(r'\d+', os.path.basename(f)).group())
)
doc = docx.Document(template_files[-1])
# Extract the bullet numId from existing List Paragraph items before clearing
template_num_id = 1  # fallback
for p in doc.paragraphs:
    if p.style.name == 'List Paragraph':
        numPr = p._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
        if numPr is not None:
            template_num_id = int(numPr.get(qn('w:val')))
            break
# Clear all existing content but preserve sectPr (page margins/layout)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child is not sectPr:
        body.remove(child)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_normal(text):
    doc.add_paragraph(text, style='Normal')

def add_centered(text):
    p = doc.add_paragraph(text, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading_underline(text):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.underline = True

def add_list_item(text):
    """Add a List Paragraph with bullet numbering from the template."""
    p = doc.add_paragraph(text, style='List Paragraph')
    # Don't set left_indent explicitly -- let the numbering definition handle it
    # (146 uses left=1444 twips ~1" with hanging=360 ~0.25")
    # Explicitly add bullet numbering (numId from template's List Paragraph)
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(template_num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def add_hyperlink_run(paragraph, url, text=None):
    """Add a clickable hyperlink run to an existing paragraph."""
    if text is None:
        text = url
    r_id = doc.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_link(tabs, text):
    """Add a line with URL formatted as a clickable blue hyperlink."""
    import re
    p = doc.add_paragraph(style='Normal')
    if tabs:
        p.add_run(tabs)
    url_match = re.search(r'(https?://\S+|www\.\S+)', text)
    if url_match:
        before = text[:url_match.start()]
        url_text = url_match.group(1)
        after = text[url_match.end():]
        if before:
            p.add_run(before)
        add_hyperlink_run(p, url_text if url_text.startswith('http') else 'https://' + url_text, url_text)
        if after:
            p.add_run(after)
    else:
        p.add_run(text)

def add_topic(title, body):
    p = doc.add_paragraph(style='Normal')
    p.add_run('\t')
    run_title = p.add_run(f'{title}: ')
    run_title.bold = True
    p.add_run(body)

# === SECTION 1: Title ===
add_centered('Consultation')

# === SECTION 2: Project Plan ===
add_heading_underline('Project plan')
add_normal('\tThe client is subscribed to an automated trading service that executes credit spreads on index options (SPX, NDX) in a master Tradier account. The client wants a trade copier bot that monitors the master account for new orders and automatically forwards them to up to 6 follower Tradier accounts. Each follower account has a configurable position size multiplier (1x, 2x, 3x, etc.). The bot must preserve multi-leg order structure (credit spreads cannot be broken into individual legs), prevent duplicate order processing, and handle stale/unfilled orders. The bot will be deployed to the cloud (Heroku) with a web-based dashboard for managing accounts and settings.')

# === SECTION 3: Time and Cost Estimate ===
add_heading_underline('Time and cost estimate')
add_normal('\tSee Excel sheet for breakdown.')
add_normal('\tTime: 20 hours @ 2 hours per day \u2192 10 days to delivery')
add_normal('\tInitial Cost: $2,000')
add_normal('\tOngoing costs: $7/month')
add_list_item('MongoDB: $0/month (would be $9 if usage ever gets very big)')
add_list_item('Heroku: $7/month')
add_list_item('Github: $0/month')
add_list_item('Tradier: $0/month')

# === SECTION 4: Meeting Summary ===
add_heading_underline('Meeting summary')
add_normal('\tThe client and I discussed several topics, including the following.')

add_topic('Question sheet', 'I sent the client a question sheet prior to the meeting and reviewed the responses. Most answers were straightforward. I needed to clarify the asset class, as the client listed "indexes and stocks" but I knew you cannot trade indexes directly. The client confirmed he meant options on the indexes (SPX, NDX), primarily credit spreads.')

add_topic('Copy trading concept', 'The client subscribes to an automated signal service that trades credit spreads in a master Tradier account. The service costs $150 per month per contract. As the client\'s account grows, scaling up requires paying proportionally more ($150 per additional contract). The client wants to avoid this cost by building a custom copy trading bot that monitors the master account and forwards trades to his other accounts. The signal service enters 0DTE credit spreads around 10:30 AM when the market begins to go flat, collecting premium that decays to zero by end of day. The client has been using the service for about a month and a half with no losing days.')

add_topic('Brokerage experience', 'The client asked about my experience with Tradier. I stated that I have used Tradier for approximately a dozen jobs. Most Tradier clients are attracted to the commission-free options trading, paying only about $10 per month for the account subscription.')

add_topic('Multi-leg orders', 'The client emphasized that multi-leg orders (credit spreads) must be forwarded as a single order and never broken into individual legs, as having one leg fill without the other could be devastating. I stated that the Tradier API returns multi-leg orders as a single entity, so the bot would not attempt to break them apart. I also noted that the broker would reject individual naked option legs due to margin requirements, providing an additional safety layer.')

add_topic('Duplicate trade prevention', 'The client asked about the risk of the bot reading an order, attempting to copy it, then reading the same order again and duplicating it. I stated that we would use a database to track each processed order ID and continuously check against it to ensure no order is sent more than once.')

add_topic('Order monitoring and latency', 'The client asked whether the bot would use webhooks. I stated that no webhooks are involved; the bot continuously polls the Tradier API for new orders and detects new ones as they appear. The client asked about potential lag. I stated there would be approximately one second of delay between master execution and follower order placement, which the client found acceptable. I also noted that if the master account were to place an order one second before market close, the followers might not get it in time, but the client confirmed all orders come through between 10:30 AM and 12:00 PM.')

add_topic('Stale orders and order expiration', 'The client described a scenario where a limit order sat unfilled until 3:00 PM and he was unsure what to do. I stated that the bot could track how long each order has been on the market and cancel it after a configurable timeout. The client noted that with 0DTE options, anything less than 2 hours before close has very little upside. I also asked about early market close days, and the client agreed the bot should be aware of those as well.')

add_topic('Position sizing', 'The client stated that the master account trades one contract at a time, but the client wants to scale up as his account grows (e.g., 1x, 2x, 3x the master position). I confirmed we could build in a configurable multiplier per follower account.')

add_topic('Follower accounts', 'The client stated he wants to forward trades to up to 6 follower accounts, including accounts in a spouse\'s name. I stated that from a software perspective, the bot can forward to as many accounts as the client has API credentials for. The client confirmed he has spoken with Tradier tech support about getting written authorization to trade on the spouse\'s account. I noted that there may be legal implications I cannot advise on, as I am not a lawyer.')

add_topic('Pattern day trading rule', 'I mentioned that FINRA has formally sent a request to the SEC for margin reform and the PDT rule\'s $25,000 requirement might be eliminated, possibly as soon as April. The client was interested, noting that the signal service would probably execute more trades per day if the rule were removed. I also mentioned that some traders currently let 0DTE positions expire rather than closing them to avoid triggering day trades.')

add_topic('Competitor services', 'I mentioned Traders Connect (https://tradersconnect.com/) as a copy trading platform, though it likely does not support Tradier. The client stated that all existing copy trading services he found required linking to a licensed professional rather than copying from another personal account. The client said Tradier tech support confirmed it could be done via the API but they do not have it built out natively, which is why the client sought a custom solution.')

add_topic('Code hosting and cloud deployment', 'I stated that the code would be hosted on GitHub in the client\'s name. The client would create a GitHub account and add me as a collaborator. This ensures the client retains ownership and access even if another developer is brought on later. From GitHub, the code can be connected to a cloud hosting platform such as Heroku for deployment.')

add_topic('Commissions', 'The client noted that closing commissions on some trades could be $70 to $80, which significantly impacts profit on a $5,000 account making $150 per day. This is part of why the signal service uses 0DTE credit spreads that expire worthless rather than being actively closed.')

add_topic('Timeline', 'I stated that for a trade copier, the project would take approximately two weeks. I could have something deliverable within one week and finalized within two weeks.')

add_topic('Pricing', 'I stated that my average job costs around $3,000, but this project would likely come in below that since there is no market data monitoring or indicator calculation involved. I estimated approximately $2,000.')

add_topic('Next steps', 'I stated that within 24 hours I would return an estimate for client review. If the client wants to open a contract, he will need to decide between fixed price and hourly contract type. I offered to include a setup guide and optionally do a remote walkthrough session via AnyDesk after delivery. The client expressed interest in approving the estimate the same day.')

# === SECTION 5: Helpful Links (BOILERPLATE) ===
# Pattern from 146: heading underlined, categories at 1 tab, URLs at 2 tabs
L = '\t'    # 1 tab - category headers
U = '\t\t'  # 2 tabs - URLs

add_heading_underline('Helpful links:')
add_normal(f'{L}My website')
add_link(U, 'https://tnttrading.net/')

add_normal(f'{L}Web-Based Automated Trading Platforms:')
for url in [
    'https://traderspost.io/', 'https://whispertrades.com/',
    'https://www.tradesteward.com/', 'https://signalstack.com/',
    'https://capitalise.ai/', 'https://www.machinetrader.io/',
    'https://www.quantconnect.com/', 'https://www.quantrocket.com/',
    'https://www.composer.trade/ (includes pre-made bots)',
    'https://www.surmount.ai/ (includes pre-made bots)',
    'https://optionalpha.com/ (includes pre-made bots)',
    'https://crosstrade.io/',
]:
    add_link(U, url)

add_normal(f'{L}Copy Trading Services:')
for url in [
    'https://trade.collective2.com/', 'https://tradersconnect.com/',
    'https://www.quiverquant.com/', 'https://www.getquantbase.com/',
    'https://www.etoro.com/', 'https://www.zulutrade.com/',
    'https://www.darwinex.com/', 'https://nagamarkets.com/',
]:
    add_link(U, url)

add_normal(f'{L}Webhooks:')
add_link(U, 'https://www.tradingview.com/ (includes pre-made scripts)')
add_link(U, 'https://trendspider.com/')

add_normal(f'{L}Equity backtesting (no-code):')
add_link(U, 'https://finviz.com/')

add_normal(f'{L}Options backtesting (no-code):')
for url in [
    'https://orats.com/backtester', 'https://optionomega.com/',
    'https://www.edeltapro.com/', 'https://wallstreet.io/',
    'https://www.optioncolors.com/', 'https://www.trademachine.com/',
    'https://tradeautomationtoolbox.com/',
]:
    add_link(U, url)

add_normal(f'{L}Options strategy calculator:')
add_link(U, 'https://www.macroption.com/calculators/')

add_normal(f'{L}Brokerage Account Linking Services:')
for url in [
    'https://snaptrade.com/', 'https://www.meshconnect.com/',
    'https://plaid.com/',
]:
    add_link(U, url)

add_normal(f'{L}Equity data providers:')
for url in [
    'https://twelvedata.com/', 'https://finnhub.io/',
    'https://www.alphavantage.co/', 'https://www.tiingo.com/',
    'https://site.financialmodelingprep.com/', 'https://eodhd.com/',
]:
    add_link(U, url)

add_normal(f'{L}Options data providers:')
for url in [
    'https://polygon.io/', 'https://www.activetick.com/',
    'https://www.thetadata.net/',
    'https://intrinio.com/options/options-realtime',
    'https://www.ivolatility.com/landing/data-cloud-api',
    'https://www.optiondata.io/', 'https://tradefeeds.com/',
    'https://finnworlds.com/options-chain-api/',
]:
    add_link(U, url)

add_normal(f'{L}Futures data providers:')
add_link(U, 'https://databento.com/')

add_normal(f'{L}Cloud platforms:')
for url in [
    'https://www.heroku.com/', 'https://fly.io/', 'https://render.com/',
]:
    add_link(U, url)

add_normal(f'{L}Databases:')
for url in [
    'https://www.mongodb.com/',
    'https://firebase.google.com/products/realtime-database',
    'https://deta.space/',
]:
    add_link(U, url)

add_normal(f'{L}Web frameworks:')
for url in [
    'https://flask.palletsprojects.com/', 'https://plotly.com/dash/',
    'https://streamlit.io/',
]:
    add_link(U, url)

add_normal(f'{L}Algo-Trading Communities:')
for url in [
    'https://www.reddit.com/r/algotrading/ (1.7 million members)',
    'https://discord.com/invite/BEr6y6Xqyv (for TD Ameritrade Python API, 3800 members)',
    'https://discord.com/invite/TmMsJCKY3T (for LumiBot, 2200 members)',
    'https://discord.com/invite/tnhrAfhNC9 (for TraderOracle, 1300 members)',
    'https://discord.com/invite/shPCaPcUr9 (for TraderMisfits, 145 members)',
]:
    add_link(U, url)

add_normal(f'{L}YouTube channels:')
for url in [
    'https://www.youtube.com/@Algovibes (119k subs)',
    'https://www.youtube.com/@parttimelarry (108k subs)',
    'https://www.youtube.com/@ChadThackray (17k subs)',
]:
    add_link(U, url)

add_normal(f'{L}Udemy courses:')
for url in [
    'https://www.udemy.com/course/python-for-finance-and-trading-algorithms/',
    'www.udemy.com/course/algorithmic-trading-quantitative-analysis-using-python/',
    'www.udemy.com/course/algorithmic-trading-with-python-and-machine-learning/',
]:
    add_link(U, url)

add_normal(f'{L}Academic articles related to trading:')
for url in [
    'VWAP: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4631351',
    'ORB: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4729284',
    'Fill execution: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4189239',
    'Profitability: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4416622',
    'Martingale: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4678427',
    'Momentum: https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4824172',
]:
    add_link(U, url)

add_normal(f'{L}Funded account services:')
for url in [
    'https://www.topstep.com/', 'https://tradeday.com/',
    'https://ftmo.com/en/', 'https://apextraderfunding.com/',
    'https://myfundedfutures.com/', 'https://shop.fasttracktrading.net/',
]:
    add_link(U, url)

add_normal(f'{L}Agencies offering trading bot automation services:')
for url in [
    'https://greyhoundanalytics.com/', 'https://affordableindicators.com/',
    'https://www.pika.group/', 'https://polygant.net/',
    'https://www.effectivesoft.com/automated-trading-software-development.html',
    'https://www.fourchaintech.com/cryptocurrency-trading-bot-development',
    'https://itexus.com/top-10-stock-trading-bot-developers/',
]:
    add_link(U, url)

# === SECTION 6: My Own Websites ===
add_normal(f'{L}My own websites (these are password protected, but I have some screenshots below)')
for url in [
    'https://www.archpublic.pro/',
    'https://tb-ic2486-d42392745503.herokuapp.com/',
    'https://option-trading-bot-149d5d0063d2.herokuapp.com/',
    'https://jobi-ustarde-app-9847db156bf0.herokuapp.com/',
    'https://stikeleathertrading-5e2f2bcd916e.herokuapp.com/',
]:
    add_link(U, url)

add_normal('')  # blank line before Photos

# === SECTION 7: Photos ===
add_heading_underline('Photos:')
add_normal('Below are some screenshots of some websites I have made for some clients')
# Screenshot images extracted from most recent consultation
SCREENSHOTS_DIR = os.path.join(CONSULTATIONS_DIR, 'temp_screenshots')
caption_images = {
    'Login page': ['00_Login_page.png'],
    'Accounts page': ['01_Accounts_page.png'],
    'Activity page': ['02_Activity_page.png'],
    'Balances page': ['03_Balances_page.png', '04_Balances_page.png'],
    'Chart page': ['05_Chart_page.png', '06_Chart_page.png', '07_Chart_page.png', '08_Chart_page.png'],
    'Orders page': ['09_Orders_page.png'],
    'PnL page': ['10_PnL_page.png', '11_PnL_page.png', '12_PnL_page.png'],
    'Positions page': ['13_Positions_page.png'],
    'Settings page': ['14_Settings_page.png', '15_Settings_page.png', '16_Settings_page.png'],
    'Trade page': ['17_Trade_page.png'],
    'Universe / watchlist page': ['18_Universe___watchlist_page.png'],
    'MongoDB database setup': ['19_MongoDB_database_setup.png'],
    'Github README': ['20_Github_README.png'],
}
from docx.shared import Inches as DocInches
for caption in caption_images:
    add_normal(caption)
    for img_file in caption_images[caption]:
        img_path = os.path.join(SCREENSHOTS_DIR, img_file)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=DocInches(6.5))

# === SAVE ===
output_path = f'{CONSULTATIONS_DIR}\\Consultation{CONSULTATION_NUMBER}.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
print(f'\nTopics: 15')
print('1. Question sheet')
print('2. Copy trading concept')
print('3. Brokerage experience')
print('4. Multi-leg orders')
print('5. Duplicate trade prevention')
print('6. Order monitoring and latency')
print('7. Stale orders and order expiration')
print('8. Position sizing')
print('9. Follower accounts')
print('10. Pattern day trading rule')
print('11. Competitor services')
print('12. Code hosting and cloud deployment')
print('13. Commissions')
print('14. Timeline')
print('15. Pricing')
print('16. Next steps')
