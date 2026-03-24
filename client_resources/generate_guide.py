#!/usr/bin/env python3
"""
Setup Guide Generator - Template Script

This script is the TEMPLATE that the skill copies into client_resources/generate_guide.py.
Claude customizes the SERVICES list and CONFIG vars before writing to client_resources/.

Generates both .docx and .pdf setup guides with embedded screenshots.

Usage:
    python client_resources/generate_guide.py

Dependencies:
    pip install python-docx fpdf2
"""

import os
import json
import sys
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# ============================================================
# CONFIGURATION - Customized per project by the skill
# ============================================================

# Services to include in this guide (ordered)
SERVICES = ["tradier", "github", "heroku", "mongodb"]

# Configurable profile info
PROFILE_NAME = "Tyler Potts"
PROFILE_EMAIL = "twpotts11@gmail.com"
PROFILE_USERNAME = "twpotts"
PROFILE_URL = "upwork.com/freelancers/robotraderguy"

# Output settings
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DATE_STR = date.today().strftime('%Y-%m-%d')
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, f"SetupGuide {DATE_STR}.docx")
OUTPUT_PDF = os.path.join(OUTPUT_DIR, f"SetupGuide {DATE_STR}.pdf")

# Paths
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(PROJECT_ROOT, '.claude', 'skills',
                          'skill-setup-guide', 'assets')
LOGO_PATH = (r"C:\Users\twpot\Documents\Tyler's Stuff\Tyler's Documents"
             r"\TNT Trading Partners\Cloud\tnttrading\static\img\upwork_logo.png")

# Display names for services
DISPLAY_NAMES = {
    'github': 'Github',
    'heroku': 'Heroku',
    'mongodb': 'MongoDB',
    'tradier': 'Tradier',
    'alpaca': 'Alpaca',
    'schwab': 'Charles Schwab',
    'coinbase': 'Coinbase',
    'oanda': 'Oanda',
    'tradestation': 'TradeStation',
    'tastytrade': 'TastyTrade',
    'tradovate': 'Tradovate',
    'gmail': 'Gmail',
    'twilio': 'Twilio',
    'discord': 'Discord',
    'papertrail': 'Papertrail',
    'clicksend': 'ClickSend',
    'etrade': 'E*TRADE',
}

# Upwork green color
UPWORK_GREEN = (20, 168, 0)

# Brand colors for each service (RGB tuples)
BRAND_COLORS = {
    'github':       (36, 41, 46),      # GitHub dark charcoal
    'heroku':       (67, 0, 152),       # Heroku purple
    'mongodb':      (0, 104, 74),       # MongoDB forest green
    'tradier':      (0, 122, 204),      # Tradier blue
    'alpaca':       (252, 196, 25),     # Alpaca gold
    'schwab':       (0, 160, 223),      # Schwab blue
    'coinbase':     (0, 82, 255),       # Coinbase blue
    'oanda':        (2, 48, 71),        # Oanda dark navy
    'tradestation': (0, 51, 102),       # TradeStation navy
    'tastytrade':   (200, 16, 46),      # TastyTrade red
    'tradovate':    (30, 136, 229),     # Tradovate blue
    'gmail':        (219, 68, 55),      # Gmail red
    'twilio':       (241, 78, 50),      # Twilio red-orange
    'discord':      (88, 101, 242),     # Discord blurple
    'papertrail':   (51, 122, 183),     # Papertrail blue
    'clicksend':    (0, 166, 81),       # ClickSend green
    'etrade':       (82, 45, 128),      # E*TRADE purple
}


# ============================================================
# LOAD SERVICE CONTENT
# ============================================================

def load_service_content(service_name):
    """Load content.json for a service from the skill's assets."""
    content_path = os.path.join(ASSETS_DIR, service_name, 'content.json')
    if not os.path.exists(content_path):
        print(f"  WARNING: No content found for '{service_name}' at {content_path}")
        return None
    with open(content_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_screenshot_path(service_name, filename):
    """Get full path to a screenshot file."""
    return os.path.join(ASSETS_DIR, service_name, 'screenshots', filename)


def get_logo_path(service_name):
    """Get full path to a service logo file."""
    return os.path.join(ASSETS_DIR, 'logos', f'{service_name}.png')


# Short descriptions for cover page cards
SERVICE_DESCRIPTIONS = {
    'github': 'Cloud code storage',
    'heroku': 'Cloud hosting & deployment',
    'mongodb': 'Cloud database',
    'tradier': 'Brokerage API',
    'alpaca': 'Brokerage API',
    'schwab': 'Brokerage API',
    'coinbase': 'Crypto exchange API',
    'oanda': 'Forex brokerage API',
    'tradestation': 'Brokerage API',
    'tastytrade': 'Brokerage API',
    'tradovate': 'Futures brokerage API',
    'gmail': 'Email notifications',
    'twilio': 'SMS notifications',
    'discord': 'Chat notifications',
    'papertrail': 'Log management',
    'clicksend': 'SMS notifications',
    'etrade': 'Brokerage API',
}


# ============================================================
# DOCX GENERATION
# ============================================================

def filter_elements(elements, service_name, strip_final_note=False):
    """Filter out intro list items, duplicate section headings, and optionally final notes."""
    import re
    filtered = []
    seen_why = False
    for elem in elements:
        if elem['type'] != 'text':
            filtered.append(elem)
            continue
        text = elem['content'].strip()
        text_lower = text.lower()
        # Skip numbered intro list items like "1: Tradier (brokerage service)"
        if re.match(r'^\d+[:\.\)]\s*\w+.*\(.*service\)', text, re.IGNORECASE):
            continue
        # Skip duplicate section heading like "1: Tradier" that matches
        # the heading we already add in the generator
        if not seen_why and re.match(r'^\d+[:\.\)]\s*', text) and len(text) < 30:
            continue
        if 'why do you need' in text_lower:
            seen_why = True
        # Strip "Final note:" and everything after from the last service
        if strip_final_note and text_lower.startswith('final note'):
            break
        filtered.append(elem)
    return filtered


def extract_final_note(elements):
    """Extract 'Final note:' and subsequent text from elements."""
    final_elements = []
    capturing = False
    for elem in elements:
        if elem['type'] == 'text' and elem['content'].strip().lower().startswith('final note'):
            capturing = True
        if capturing:
            final_elements.append(elem)
    return final_elements


def generate_docx():
    """Generate the .docx setup guide."""
    from docx.oxml.ns import qn as docx_qn
    from docx.oxml import OxmlElement

    print(f"\nGenerating DOCX: {OUTPUT_DOCX}")
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)

    # === Upwork Header Banner (on all pages) ===
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Add logo to header if available
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        run = header_para.add_run()
        run.add_picture(LOGO_PATH, height=Pt(14))
        run = header_para.add_run('  ')

    # Profile name
    run = header_para.add_run(PROFILE_NAME)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 30, 0)
    run = header_para.add_run('  |  ')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run = header_para.add_run(PROFILE_URL)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Green bottom border on header paragraph
    pPr = header_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(docx_qn('w:val'), 'single')
    bottom.set(docx_qn('w:sz'), '6')
    bottom.set(docx_qn('w:space'), '1')
    bottom.set(docx_qn('w:color'), '14A800')  # Upwork green
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === Footer with page numbers ===
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(docx_qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(docx_qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(docx_qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

    # Title
    doc.add_paragraph()  # top spacing
    title = doc.add_heading(f'Setup Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(*UPWORK_GREEN)
        run.font.size = Pt(28)

    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DATE_STR)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # Intro paragraph
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run('This setup guide includes instructions for cloud deployment.')
    run.font.size = Pt(11)
    run2 = intro.add_run('\nWe will need to create the following accounts:')
    run2.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Service cards as a table with logos
    num_services = len(SERVICES)
    table = doc.add_table(rows=num_services, cols=3, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, service in enumerate(SERVICES):
        display = DISPLAY_NAMES.get(service, service.title())
        desc = SERVICE_DESCRIPTIONS.get(service, '')
        brand_clr = BRAND_COLORS.get(service, UPWORK_GREEN)
        logo_path = get_logo_path(service)

        row = table.rows[i]
        # Set row height
        row.height = Pt(45)

        # Column 0: Logo
        cell0 = row.cells[0]
        cell0.width = Pt(50)
        cell0_para = cell0.paragraphs[0]
        cell0_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(logo_path):
            run = cell0_para.add_run()
            try:
                run.add_picture(logo_path, height=Pt(28))
            except Exception:
                run.add_text(display[0])

        # Column 1: Service name
        cell1 = row.cells[1]
        cell1.width = Pt(150)
        cell1_para = cell1.paragraphs[0]
        cell1_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell1_para.add_run(f'{i+1}. {display}')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*brand_clr)

        # Column 2: Description
        cell2 = row.cells[2]
        cell2_para = cell2.paragraphs[0]
        cell2_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell2_para.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.italic = True

    doc.add_paragraph()  # spacing

    # Add each service section
    for section_num, service in enumerate(SERVICES, 1):
        content = load_service_content(service)
        if not content:
            # Add placeholder section
            doc.add_heading(DISPLAY_NAMES.get(service, service.title()), level=1)
            doc.add_paragraph(f'Instructions for {DISPLAY_NAMES.get(service, service.title())} - TBA')
            continue

        display_name = content.get('display_name', service.title())

        # Section heading with brand color
        brand_color = BRAND_COLORS.get(service, UPWORK_GREEN)
        heading = doc.add_heading(f'{section_num}: {display_name}', level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*brand_color)

        # Add brand-colored accent bar under heading
        accent_para = doc.add_paragraph()
        accent_pPr = accent_para._element.get_or_add_pPr()
        accent_pBdr = OxmlElement('w:pBdr')
        accent_top = OxmlElement('w:top')
        accent_top.set(docx_qn('w:val'), 'single')
        accent_top.set(docx_qn('w:sz'), '12')
        accent_top.set(docx_qn('w:space'), '1')
        accent_top.set(docx_qn('w:color'), '{:02X}{:02X}{:02X}'.format(*brand_color))
        accent_pBdr.append(accent_top)
        accent_pPr.append(accent_pBdr)

        # Process elements (filtered to remove intro cruft)
        is_last = (section_num == len(SERVICES))
        elements = filter_elements(content.get('elements', []), service,
                                   strip_final_note=is_last)
        for elem in elements:
            if elem['type'] == 'text':
                text = elem['content']

                # Replace placeholder tokens
                text = text.replace('{EMAIL}', PROFILE_EMAIL)
                text = text.replace('{USERNAME}', PROFILE_USERNAME)
                text = text.replace('{PROFILE_NAME}', PROFILE_NAME)

                text_lower = text.lower().strip()

                # --- "Why do you need X?" as a styled sub-heading ---
                if text_lower.startswith('why do you need'):
                    h = doc.add_heading(text, level=2)
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(*brand_color)
                        run.font.size = Pt(12)

                # --- "Link:" line - render as key: value with blue link ---
                elif text_lower.startswith('link:'):
                    p = doc.add_paragraph()
                    run = p.add_run('Link: ')
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    url = text.split(':', 1)[1].strip()
                    link_run = p.add_run(url)
                    link_run.font.size = Pt(10)
                    link_run.font.color.rgb = RGBColor(0, 102, 204)

                # --- "Tutorial:" line ---
                elif text_lower.startswith('tutorial:'):
                    p = doc.add_paragraph()
                    run = p.add_run('Tutorial: ')
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    url = text.split(':', 1)[1].strip()
                    link_run = p.add_run(url)
                    link_run.font.size = Pt(10)
                    link_run.font.color.rgb = RGBColor(0, 102, 204)

                # --- "Cost:" line ---
                elif text_lower.startswith('cost:'):
                    p = doc.add_paragraph()
                    run = p.add_run('Cost: ')
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    cost_val = text.split(':', 1)[1].strip()
                    run2 = p.add_run(cost_val)
                    run2.font.size = Pt(10)

                # --- Bare URL (standalone link, skip as duplicate) ---
                elif text_lower.startswith('http'):
                    continue  # Already rendered inline with Link:/Tutorial:

                # --- Bold text ---
                elif elem.get('bold'):
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)

                # --- List items ---
                elif elem.get('list'):
                    doc.add_paragraph(text, style='List Bullet')

                # --- Normal paragraph ---
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.size = Pt(10)

            elif elem['type'] == 'image':
                img_path = get_screenshot_path(service, elem['file'])
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f'[Image: {elem["file"]} - {e}]')
                else:
                    doc.add_paragraph(f'[Image not found: {elem["file"]}]')

        # Add page break between sections
        doc.add_page_break()

    # === Conclusion Page ===
    heading = doc.add_heading('Next Steps', level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*UPWORK_GREEN)

    doc.add_paragraph()

    # Extract and render any final note from the last service
    last_content = load_service_content(SERVICES[-1])
    if last_content:
        final_elems = extract_final_note(last_content.get('elements', []))
        for elem in final_elems:
            if elem['type'] == 'text':
                text = elem['content']
                text = text.replace('{EMAIL}', PROFILE_EMAIL)
                text = text.replace('{USERNAME}', PROFILE_USERNAME)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(10)

    doc.add_paragraph()

    # Quick reference links
    ref_heading = doc.add_heading('Quick Reference', level=2)
    for run in ref_heading.runs:
        run.font.color.rgb = RGBColor(*UPWORK_GREEN)

    for service in SERVICES:
        display = DISPLAY_NAMES.get(service, service.title())
        brand_clr = BRAND_COLORS.get(service, UPWORK_GREEN)
        svc_content = load_service_content(service)
        if not svc_content:
            continue
        # Find the Link: URL from elements
        link_url = None
        for elem in svc_content.get('elements', []):
            if elem['type'] == 'text' and elem['content'].lower().startswith('link:'):
                link_url = elem['content'].split(':', 1)[1].strip()
                break
        if link_url:
            p = doc.add_paragraph()
            run = p.add_run(f'{display}: ')
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*brand_clr)
            link_run = p.add_run(link_url)
            link_run.font.size = Pt(10)
            link_run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'Questions? Contact {PROFILE_NAME} at {PROFILE_EMAIL}')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"  Saved: {OUTPUT_DOCX}")
    print(f"  Size: {os.path.getsize(OUTPUT_DOCX) / 1024 / 1024:.1f} MB")


# ============================================================
# PDF GENERATION (with Upwork banner styling)
# ============================================================

class SetupGuidePDF(FPDF):
    """PDF with Upwork-branded header and footer."""

    _current_bg = None  # (r, g, b) tint color for current section pages

    def header(self):
        # Paint page background tint if set
        if self._current_bg:
            self.set_fill_color(*self._current_bg)
            self.rect(0, 0, self.w, self.h, 'F')

        # Upwork logo (if available)
        if os.path.exists(LOGO_PATH):
            self.image(LOGO_PATH, x=10, y=7, h=7)
            self.set_xy(19, 7)
        else:
            self.set_xy(10, 7)

        # Profile name
        self.set_font('Segoe', 'B', 9)
        self.set_text_color(0, 30, 0)
        self.cell(0, 4, PROFILE_NAME)

        # Profile URL
        if os.path.exists(LOGO_PATH):
            self.set_xy(19, 11)
        else:
            self.set_xy(10, 11)
        self.set_font('Segoe', '', 7)
        self.set_text_color(100, 100, 100)
        self.cell(0, 4, PROFILE_URL)

        # Green divider line
        self.set_draw_color(*UPWORK_GREEN)
        self.set_line_width(0.5)
        self.line(10, 16, self.w - 10, 16)

        # Reset
        self.set_text_color(0, 0, 0)
        self.set_draw_color(0, 0, 0)
        self.set_line_width(0.2)
        self.set_y(19)

    def footer(self):
        self.set_y(-15)
        self.set_font('Segoe', 'I', 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')
        self.set_text_color(0, 0, 0)


def generate_pdf():
    """Generate the .pdf setup guide with Upwork-branded styling."""
    print(f"\nGenerating PDF: {OUTPUT_PDF}")

    pdf = SetupGuidePDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Register fonts
    pdf.add_font('Segoe', '', r'C:\Windows\Fonts\segoeui.ttf')
    pdf.add_font('Segoe', 'B', r'C:\Windows\Fonts\segoeuib.ttf')
    pdf.add_font('Segoe', 'I', r'C:\Windows\Fonts\segoeuii.ttf')
    pdf.add_font('Segoe', 'BI', r'C:\Windows\Fonts\segoeuiz.ttf')

    W = pdf.epw  # effective page width

    # === Title Page ===
    pdf.add_page()

    # Title
    pdf.set_font('Segoe', 'B', 22)
    pdf.set_text_color(*UPWORK_GREEN)
    pdf.cell(0, 15, 'Setup Guide', align='C')
    pdf.ln(18)

    # Date
    pdf.set_font('Segoe', '', 12)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, DATE_STR, align='C')
    pdf.ln(12)
    pdf.set_text_color(0, 0, 0)

    # Intro
    pdf.set_font('Segoe', '', 11)
    pdf.cell(0, 7, 'This setup guide includes instructions for cloud deployment.', align='C')
    pdf.ln(6)
    pdf.cell(0, 7, 'We will need to create the following accounts:', align='C')
    pdf.ln(12)

    # Create internal link targets for each section
    section_links = {}
    for service in SERVICES:
        section_links[service] = pdf.add_link()

    # Service cards with logos (clickable)
    card_h = 18
    logo_w = 14
    margin_x = 25
    card_w = W - (margin_x * 2 - 20)

    for i, service in enumerate(SERVICES):
        display = DISPLAY_NAMES.get(service, service.title())
        desc = SERVICE_DESCRIPTIONS.get(service, '')
        brand_clr = BRAND_COLORS.get(service, UPWORK_GREEN)
        logo_path = get_logo_path(service)

        y_start = pdf.get_y()

        # Brand-colored left accent bar
        pdf.set_fill_color(*brand_clr)
        pdf.rect(margin_x, y_start, 3, card_h, 'F')

        # Light background
        pdf.set_fill_color(248, 248, 248)
        pdf.rect(margin_x + 3, y_start, card_w - 3, card_h, 'F')

        # Border
        pdf.set_draw_color(220, 220, 220)
        pdf.set_line_width(0.2)
        pdf.rect(margin_x, y_start, card_w, card_h, 'D')

        # Clickable overlay for entire card
        pdf.link(margin_x, y_start, card_w, card_h, section_links[service])

        # Logo
        logo_x = margin_x + 6
        if os.path.exists(logo_path):
            try:
                pdf.image(logo_path, x=logo_x, y=y_start + 2, h=card_h - 4)
            except Exception:
                pass

        # Service name (bold, brand color)
        text_x = logo_x + logo_w + 6
        pdf.set_xy(text_x, y_start + 2)
        pdf.set_font('Segoe', 'B', 12)
        pdf.set_text_color(*brand_clr)
        pdf.cell(0, 7, f'{i+1}. {display}')

        # Description (italic, gray)
        pdf.set_xy(text_x, y_start + 9)
        pdf.set_font('Segoe', 'I', 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, desc)

        pdf.set_text_color(0, 0, 0)
        pdf.set_y(y_start + card_h + 3)

    pdf.ln(5)

    # === Service Sections ===
    for section_num, service in enumerate(SERVICES, 1):
        content = load_service_content(service)
        display_name = DISPLAY_NAMES.get(service, service.title())

        # Set page background tint for this section (applied by header on every page)
        brand_color = BRAND_COLORS.get(service, UPWORK_GREEN)
        bg_r = int(brand_color[0] * 0.08 + 255 * 0.92)
        bg_g = int(brand_color[1] * 0.08 + 255 * 0.92)
        bg_b = int(brand_color[2] * 0.08 + 255 * 0.92)
        pdf._current_bg = (bg_r, bg_g, bg_b)

        pdf.add_page()

        # Set link target so cover page cards jump here
        pdf.set_link(section_links[service])

        # Section banner card (same style as cover page)
        desc = SERVICE_DESCRIPTIONS.get(service, '')
        logo_path = get_logo_path(service)

        banner_h = 22
        banner_x = 10
        banner_w = W

        y_start = pdf.get_y()

        # Brand-colored left accent bar
        pdf.set_fill_color(*brand_color)
        pdf.rect(banner_x, y_start, 4, banner_h, 'F')

        # Light background
        pdf.set_fill_color(245, 245, 245)
        pdf.rect(banner_x + 4, y_start, banner_w - 4, banner_h, 'F')

        # Border
        pdf.set_draw_color(200, 200, 200)
        pdf.set_line_width(0.3)
        pdf.rect(banner_x, y_start, banner_w, banner_h, 'D')

        # Logo
        logo_x = banner_x + 8
        if os.path.exists(logo_path):
            try:
                pdf.image(logo_path, x=logo_x, y=y_start + 2, h=banner_h - 4)
            except Exception:
                pass

        # Service name
        text_x = logo_x + 20
        pdf.set_xy(text_x, y_start + 3)
        pdf.set_font('Segoe', 'B', 16)
        pdf.set_text_color(*brand_color)
        pdf.cell(0, 8, f'{section_num}: {display_name}')

        # Description
        pdf.set_xy(text_x, y_start + 12)
        pdf.set_font('Segoe', 'I', 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, desc)

        pdf.set_text_color(0, 0, 0)
        pdf.set_draw_color(0, 0, 0)
        pdf.set_y(y_start + banner_h + 6)

        if not content:
            pdf.set_font('Segoe', 'I', 10)
            pdf.cell(0, 7, f'Instructions for {display_name} - TBA')
            pdf.ln()
            continue

        # Process elements (filtered to remove intro cruft)
        is_last = (section_num == len(SERVICES))
        for elem in filter_elements(content.get('elements', []), service,
                                    strip_final_note=is_last):
            if elem['type'] == 'text':
                text = elem['content']
                text = text.replace('{EMAIL}', PROFILE_EMAIL)
                text = text.replace('{USERNAME}', PROFILE_USERNAME)
                text = text.replace('{PROFILE_NAME}', PROFILE_NAME)
                text_lower = text.lower().strip()

                # Check if we need a new page
                if pdf.get_y() > pdf.h - 30:
                    pdf.add_page()

                # --- "Why do you need X?" sub-heading ---
                if text_lower.startswith('why do you need'):
                    pdf.ln(2)
                    pdf.set_font('Segoe', 'B', 12)
                    pdf.set_text_color(*brand_color)
                    pdf.multi_cell(0, 7, text)
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(2)

                # --- "Link:" line ---
                elif text_lower.startswith('link:'):
                    url = text.split(':', 1)[1].strip()
                    pdf.set_font('Segoe', 'B', 10)
                    pdf.set_text_color(80, 80, 80)
                    lbl_w = pdf.get_string_width('Link: ') + 2
                    pdf.cell(lbl_w, 6, 'Link: ')
                    pdf.set_font('Segoe', '', 10)
                    pdf.set_text_color(0, 102, 204)
                    pdf.cell(0, 6, url)
                    pdf.ln(7)
                    pdf.set_text_color(0, 0, 0)

                # --- "Tutorial:" line ---
                elif text_lower.startswith('tutorial:'):
                    url = text.split(':', 1)[1].strip()
                    pdf.set_font('Segoe', 'B', 10)
                    pdf.set_text_color(80, 80, 80)
                    lbl_w = pdf.get_string_width('Tutorial: ') + 2
                    pdf.cell(lbl_w, 6, 'Tutorial: ')
                    pdf.set_font('Segoe', '', 10)
                    pdf.set_text_color(0, 102, 204)
                    pdf.cell(0, 6, url)
                    pdf.ln(7)
                    pdf.set_text_color(0, 0, 0)

                # --- "Cost:" line ---
                elif text_lower.startswith('cost:'):
                    cost_val = text.split(':', 1)[1].strip()
                    pdf.set_font('Segoe', 'B', 10)
                    pdf.set_text_color(80, 80, 80)
                    lbl_w = pdf.get_string_width('Cost: ') + 2
                    pdf.cell(lbl_w, 6, 'Cost: ')
                    pdf.set_font('Segoe', '', 10)
                    pdf.set_text_color(0, 0, 0)
                    pdf.cell(0, 6, cost_val)
                    pdf.ln(8)

                # --- Bare URL (skip - already rendered inline) ---
                elif text_lower.startswith('http'):
                    continue

                # --- Bold ---
                elif elem.get('bold'):
                    pdf.set_font('Segoe', 'B', 10)
                    pdf.multi_cell(0, 6, text)
                    pdf.ln(2)

                # --- List ---
                elif elem.get('list'):
                    pdf.set_font('Segoe', '', 10)
                    pdf.cell(8, 6, '')
                    pdf.multi_cell(W - 8, 6, f'  {text}')
                    pdf.ln(1)

                # --- Normal ---
                else:
                    pdf.set_font('Segoe', '', 10)
                    pdf.multi_cell(0, 6, text)
                    pdf.ln(2)

            elif elem['type'] == 'image':
                img_path = get_screenshot_path(service, elem['file'])
                if os.path.exists(img_path):
                    # Skip non-standard image formats
                    ext = os.path.splitext(elem['file'])[1].lower()
                    if ext in ('.emf', '.wmf'):
                        pdf.set_font('Segoe', 'I', 8)
                        pdf.set_text_color(150, 150, 150)
                        pdf.cell(0, 5, f'[Image: {elem["file"]} - format not supported in PDF]')
                        pdf.ln()
                        pdf.set_text_color(0, 0, 0)
                        continue

                    try:
                        # Check if we have enough space
                        if pdf.get_y() > pdf.h - 80:
                            pdf.add_page()

                        # Add image centered, scaled to page width
                        img_w = W * 0.85
                        x_pos = (pdf.w - img_w) / 2
                        pdf.image(img_path, x=x_pos, w=img_w)
                        pdf.ln(4)
                    except Exception as e:
                        pdf.set_font('Segoe', 'I', 8)
                        pdf.set_text_color(150, 150, 150)
                        pdf.cell(0, 5, f'[Image error: {elem["file"]} - {e}]')
                        pdf.ln()
                        pdf.set_text_color(0, 0, 0)

    # === Conclusion Page (no background tint) ===
    pdf._current_bg = None
    pdf.add_page()

    # Heading
    pdf.set_font('Segoe', 'B', 18)
    pdf.set_text_color(*UPWORK_GREEN)
    pdf.cell(0, 12, 'Next Steps', align='C')
    pdf.ln(10)
    pdf.set_draw_color(*UPWORK_GREEN)
    pdf.set_line_width(0.5)
    pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)

    # Extract and render any final note from the last service
    last_content = load_service_content(SERVICES[-1])
    if last_content:
        final_elems = extract_final_note(last_content.get('elements', []))
        for elem in final_elems:
            if elem['type'] == 'text':
                text = elem['content']
                text = text.replace('{EMAIL}', PROFILE_EMAIL)
                text = text.replace('{USERNAME}', PROFILE_USERNAME)
                pdf.set_font('Segoe', '', 10)
                pdf.multi_cell(0, 6, text)
                pdf.ln(3)

    pdf.ln(5)

    # Quick Reference section with links
    pdf.set_font('Segoe', 'B', 14)
    pdf.set_text_color(*UPWORK_GREEN)
    pdf.cell(0, 10, 'Quick Reference')
    pdf.ln(10)
    pdf.set_text_color(0, 0, 0)

    for service in SERVICES:
        display = DISPLAY_NAMES.get(service, service.title())
        brand_clr = BRAND_COLORS.get(service, UPWORK_GREEN)
        svc_content = load_service_content(service)
        if not svc_content:
            continue
        link_url = None
        for elem in svc_content.get('elements', []):
            if elem['type'] == 'text' and elem['content'].lower().startswith('link:'):
                link_url = elem['content'].split(':', 1)[1].strip()
                break
        if link_url:
            pdf.set_font('Segoe', 'B', 10)
            pdf.set_text_color(*brand_clr)
            lbl_w = pdf.get_string_width(f'{display}: ') + 2
            pdf.cell(lbl_w, 7, f'{display}: ')
            pdf.set_font('Segoe', '', 10)
            pdf.set_text_color(0, 102, 204)
            pdf.cell(0, 7, link_url)
            pdf.ln(8)

    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    # Contact line
    pdf.set_font('Segoe', 'I', 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 7, f'Questions? Contact {PROFILE_NAME} at {PROFILE_EMAIL}', align='C')
    pdf.set_text_color(0, 0, 0)

    # Save
    pdf.output(OUTPUT_PDF)
    print(f"  Saved: {OUTPUT_PDF}")
    print(f"  Size: {os.path.getsize(OUTPUT_PDF) / 1024 / 1024:.1f} MB")


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    print("=" * 60)
    print("Setup Guide Generator")
    print("=" * 60)
    print(f"Date: {DATE_STR}")
    print(f"Services: {', '.join(DISPLAY_NAMES.get(s, s) for s in SERVICES)}")
    print(f"Assets: {ASSETS_DIR}")

    # Verify assets exist
    for service in SERVICES:
        content_path = os.path.join(ASSETS_DIR, service, 'content.json')
        if os.path.exists(content_path):
            with open(content_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            print(f"  {service}: {data['text_count']} text, {data['image_count']} images")
        else:
            print(f"  {service}: NO CONTENT FOUND")

    # Generate both formats
    generate_docx()
    generate_pdf()

    print(f"\nDone! Files saved to: {OUTPUT_DIR}")
